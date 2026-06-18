from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("docs")
OUTPUT_PATH = OUTPUT_DIR / "kids_story_agent_project_summary.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = width_to_inches(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.type = "dxa"
            tc_w.w = widths[index]


def width_to_inches(dxa):
    return Inches(dxa / 1440)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Kids Daily Story Agent Project Summary")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = False

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(
        "A concise overview of the AI agent, build workflow, prompts, iterations, and lessons learned."
    )
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.text = header
        set_cell_shading(cell, "F8F9FA")
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.bold = True

    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cell = cells[index]
            cell.text = value
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(6)
    return table


def build_document():
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("Project Overview", level=1)
    add_paragraph(
        doc,
        "I built an AI agent that helps parents and teachers create personalized kids story packages in a Streamlit web app. The agent generates a short, age-appropriate story, creates an illustration, builds a PDF, optionally creates MP3 narration, waits for human approval, and then sends the package by email.",
    )
    add_paragraph(
        doc,
        "The project is designed around a human-in-the-loop workflow: the agent can automate the heavy lifting, but the parent or teacher still reviews the story before it is delivered.",
    )

    doc.add_heading("Agent Goal", level=1)
    add_paragraph(
        doc,
        "Help parents and teachers generate, review, approve, and send personalized kids story packages in under two minutes with reliable delivery.",
    )

    doc.add_heading("Main Capabilities", level=1)
    add_bullets(
        doc,
        [
            "Collects child profile details such as name, age, gender, interests, favorite character types, and topics to avoid.",
            "Generates a short, safe, child-friendly story using the selected theme and stored preferences.",
            "Creates a story illustration using OpenAI image generation and makes the image mandatory before sending.",
            "Creates a PDF containing the story and illustration, then attaches it to the email.",
            "Optionally creates an ElevenLabs MP3 narration when the parent or teacher selects audio mode.",
            "Sends the approved story package through Gmail SMTP and protects against duplicate daily sends.",
        ],
    )

    doc.add_heading("Tools and Services Used", level=1)
    add_table(
        doc,
        ["Tool / Service", "Role in the Agent"],
        [
            ("Nebius", "Used as the main LLM provider for story generation and agent reasoning."),
            ("OpenAI Images", "Used to generate story illustrations from profile-aware prompts."),
            ("ElevenLabs", "Used to generate optional MP3 audio narration."),
            ("Gmail SMTP", "Used to send the approved story email to the parent or teacher."),
            ("Mem0", "Used for cross-session memory of child preferences and interests."),
            ("Pinecone", "Planned or optional semantic search layer for past stories and preference retrieval."),
            ("Local file storage", "Used in development for profiles, logs, generated PDFs, images, and MP3 files."),
            ("Streamlit", "Used as the web surface for profile setup, story preview, and approval."),
        ],
        [2300, 7060],
    )

    doc.add_heading("Datasets Used", level=1)
    add_paragraph(
        doc,
        "This project does not rely on a large external training dataset. It uses user-provided child profile data, generated story history, memory records, and local delivery logs as the working data for personalization.",
    )
    add_bullets(
        doc,
        [
            "Child profiles: name, age, gender, interests, favorite character types, avoided topics, and email settings.",
            "Story history: generated titles, themes, story text, parent notes, and delivery outcomes.",
            "Memory data: saved child preferences and interests through Mem0, keyed by child ID.",
            "Generated artifacts: PDFs, images, and MP3 files saved locally during development.",
            "Email logs: send status, duplicate-send checks, attachment details, and failure reasons.",
        ],
    )

    doc.add_heading("Prompts Used During Vibe Coding", level=1)
    add_paragraph(
        doc,
        "The build evolved through natural language prompts that gradually added features, fixed bugs, and improved the user experience. The most important prompts were:",
    )
    add_numbered(
        doc,
        [
            "Build an AI agent that helps parents and teachers generate personalized kids stories with a web UI.",
            "Use a child profile with name, age, gender, interests, favorite character types, topics to avoid, email, and preferred delivery time.",
            "Make the story shorter and age-appropriate.",
            "Send the email daily at 8:30 PM and prevent duplicate emails.",
            "Improve the UI because it looks too plain and all white.",
            "Take favorite character types into consideration, including Elsa-like preferences without using copyrighted characters directly.",
            "Add the story as a PDF attachment and keep the image and story on the same page when possible.",
            "Add Mem0 memory and use a common child ID across memory systems.",
            "Add a gender field and consider gender when generating images.",
            "Add optional ElevenLabs voice mode so parents can choose whether they want audio narration.",
            "If audio is requested, generate an MP3 and attach it to the email; if it fails, do not send a partial email.",
            "Make image generation mandatory before sending.",
            "Never print API keys or secrets in logs, terminal output, UI, or error messages.",
        ],
    )

    doc.add_heading("Iterations Tried", level=1)
    add_table(
        doc,
        ["Iteration", "What Changed"],
        [
            ("Email setup", "Moved from app-only preview to real Gmail SMTP delivery with environment configuration."),
            ("Story length", "Adjusted generation logic so stories are shorter and easier for kids to consume."),
            ("Scheduling", "Added a daily 8:30 PM job and duplicate-send protection."),
            ("UI design", "Simplified the app, removed unnecessary sections, and made the layout more kid-friendly."),
            ("Personalization", "Added favorite character types, memory, gender, and child-specific preferences."),
            ("Attachments", "Added PDF generation, inline image handling, and reduced email image size."),
            ("Audio mode", "Added an optional ElevenLabs toggle, MP3 generation, and attachment handling."),
            ("Reliability", "Made image mandatory, blocked sends when requested audio fails, and improved error messages."),
        ],
        [2200, 7160],
    )

    doc.add_heading("Human-in-the-Loop Design", level=1)
    add_paragraph(
        doc,
        "The agent stops before final delivery and asks the parent or teacher to approve the generated story package. The human can approve and send, revise the story, skip the email, or adjust profile settings before delivery.",
    )

    doc.add_heading("Failure Handling", level=1)
    add_bullets(
        doc,
        [
            "If image generation fails, the email is not sent because the image is mandatory.",
            "If audio narration is requested but no MP3 is generated, the email is not sent.",
            "If email delivery fails, the app records the failure reason in the local email log.",
            "If duplicate-send protection detects that the child already received today's story, the agent avoids sending another one.",
        ],
    )

    doc.add_heading("Learnings and Observations", level=1)
    add_bullets(
        doc,
        [
            "Human approval is important because the output is for children and should be reviewed before delivery.",
            "Feature toggles matter: optional audio should only be generated and attached when the parent or teacher requests it.",
            "Production storage should move from local files to cloud object storage plus a database.",
            "Pinecone is better for semantic recall and personalization than for storing files like PDFs or MP3s.",
            "Secret handling must be strict: API keys should never be printed, logged, or shown in UI output.",
            "Long-running Streamlit apps may keep stale environment values, so clean restarts are important after changing `.env`.",
            "Good agent design needs clear stop conditions, especially when required tools fail.",
        ],
    )

    doc.add_heading("Production-Ready Direction", level=1)
    add_paragraph(
        doc,
        "For production, local file storage should be replaced with cloud object storage such as AWS S3, Google Cloud Storage, Azure Blob Storage, or Cloudflare R2. Structured records such as profiles, story metadata, delivery status, and settings should move to a production database like PostgreSQL.",
    )
    add_paragraph(
        doc,
        "Email delivery should also move from Gmail SMTP to a production email provider such as SendGrid, Resend, Postmark, or Amazon SES.",
    )
    add_paragraph(
        doc,
        "In a production version, the agent can also be changed to run automatically every day at a specific time, such as a parent or teacher's preferred delivery time, using a managed scheduler or background job service.",
    )

    doc.add_heading("Success Metric", level=1)
    add_paragraph(
        doc,
        "The project is successful when parents and teachers can generate, review, approve, and send a complete story package with PDF, image, and optional audio in under two minutes with at least a 95% successful delivery rate.",
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
